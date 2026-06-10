from fastapi import FastAPI, HTTPException, UploadFile, File, Form
import pdfplumber
from docx import Document
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
import json
import uuid
import os
import shutil
from dotenv import load_dotenv
from detector import detect_columns_with_llm, build_detection_result
from database import init_db, save_mapping, get_mapping, save_upload, get_uploads
from cleaner import clean_dataframe
load_dotenv()

app = FastAPI()

# Initialize database tables when the app starts
@app.on_event("startup")
async def startup_event():
    init_db()

# CORS Middleware, allows frontend to communicate with backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory where uploaded files will be stored
UPLOAD_DIR = "uploads"
# Supported file extensions for upload
ALLOWED_EXTENSIONS = {"csv", "xlsx", "xls", "pdf", "docx"}

# Extract file extension from filename
def get_extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

# Extract tables from PDF using pdfplumber, falls back to text extraction if no tables found
def extract_pdf(file_path: str):
    tables = []
    full_text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables()
            for table in page_tables:
                if table:
                    headers = table[0]
                    rows = table[1:]
                    df = pd.DataFrame(rows, columns=headers)
                    tables.append(df)
            full_text += page.extract_text() or ""
    if tables:
        return pd.concat(tables, ignore_index=True), "table"
    lines = [line.strip() for line in full_text.splitlines() if line.strip()]
    if lines:
        return pd.DataFrame({"raw_text": lines}), "text"
    return None, None

# Extract tables from DOCX using python-docx, falls back to paragraphs if no tables found
def extract_docx(file_path: str):
    doc = Document(file_path)
    tables = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            rows.append([cell.text.strip() for cell in row.cells])
        df = pd.DataFrame(rows, columns=headers)
        tables.append(df)
    if tables:
        return pd.concat(tables, ignore_index=True), "table"
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if lines:
        return pd.DataFrame({"raw_text": lines}), "text"
    return None, None

# Read any supported file and return a dataframe
def read_file_to_df(save_path: str, ext: str):
    if ext == "csv":
        return pd.read_csv(save_path)
    elif ext in ["xlsx", "xls"]:
        return pd.read_excel(save_path)
    elif ext == "pdf":
        df, _ = extract_pdf(save_path)
        return df
    elif ext == "docx":
        df, _ = extract_docx(save_path)
        return df
    return None

# Helper function to calculate fill rate per column. Fill_rate = percentage of rows that have a value (0.0 to 1.0)
def calculate_fill_rates(df: pd.DataFrame) -> dict:
    fill_rates = {}
    total = len(df)
    for col in df.columns:
        filled = df[col].replace("", float("nan")).dropna().count()
        fill_rates[col] = round(filled / total, 2) if total > 0 else 0.0
    return fill_rates

# Root endpoint to test whether the API is running
@app.get("/")
def root():
    return {"message": "Audit AI API is running"}

# Endpoint for uploading files
@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    client_id: str = Form(...)
):
    # Validate file extension
    ext = get_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=f"File type .{ext} not supported. Upload Excel, CSV, PDF or DOCX file only."
        )
    # Check file size (max 50MB)
    MAX_FILE_SIZE = 50
    file_bytes = await file.read()
    file_size_mb = len(file_bytes) / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File size exceeds the maximum limit of {MAX_FILE_SIZE} MB. Uploaded file size: {file_size_mb:.2f} MB."
        )
    # Reset file pointer after reading for size check
    file.file.seek(0)
    # Save file to uploads directory
    file_id = str(uuid.uuid4())
    save_path = os.path.join(UPLOAD_DIR, f"{file_id}.{ext}")
    with open(save_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # PDF
    if ext == "pdf":
        df, source = extract_pdf(save_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not extract any content from PDF.")
        save_upload(file_id, client_id, file.filename, ext, len(df))
        # Calculate fill rates for column detection context
        fill_rates = calculate_fill_rates(df)
        return {
            "file_id": file_id,
            "client_id": client_id,
            "filename": file.filename,
            "source": source,
            "rows": len(df),
            "columns": list(df.columns),
            "fill_rates": fill_rates,
            "preview": df.head(5).fillna("").to_dict(orient="records"),
            "message": f"PDF uploaded — extracted via {source}"
        }

    # DOCX
    if ext == "docx":
        df, source = extract_docx(save_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not extract any content from DOCX.")
        save_upload(file_id, client_id, file.filename, ext, len(df))
        # Calculate fill rates for column detection context
        fill_rates = calculate_fill_rates(df)
        return {
            "file_id": file_id,
            "client_id": client_id,
            "filename": file.filename,
            "source": source,
            "rows": len(df),
            "columns": list(df.columns),
            "fill_rates": fill_rates,
            "preview": df.head(5).fillna("").to_dict(orient="records"),
            "message": f"DOCX uploaded — extracted via {source}"
        }

    # CSV / EXCEL
    try:
        if ext == "csv":
            df = pd.read_csv(save_path)
        else:
            df = pd.read_excel(save_path)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")
    save_upload(file_id, client_id, file.filename, ext, len(df))
    # Calculate fill rates for column detection context
    fill_rates = calculate_fill_rates(df)
    return {
        "file_id": file_id,
        "client_id": client_id,
        "filename": file.filename,
        "source": "table",
        "rows": len(df),
        "columns": list(df.columns),
        "fill_rates": fill_rates,
        "preview": df.head(5).fillna("").to_dict(orient="records"),
        "message": "File uploaded and processed successfully"
    }

# Endpoint for detecting column meanings.First checks if client has a saved mapping, if yes skips LLM entirely
# If no saved mapping found, runs LLM detection using column names, sample values and fill rates
@app.post("/detect-columns")
async def detect_columns_endpoint(
    client_id: str = Form(...),
    file_id: str = Form(...),
    columns: str = Form(...),
    file_type: str = Form("general"),
    fill_rates: str = Form("{}")  
):
    # Parse columns list from JSON string
    try:
        columns_list = json.loads(columns)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid columns format.")
    # Parse fill rates from JSON string
    try:
        fill_rates_dict = json.loads(fill_rates)
    except json.JSONDecodeError:
        fill_rates_dict = {}

    # Check if this client already has a saved mapping in the database
    saved_mapping = get_mapping(client_id, file_type)
    if saved_mapping:
        all_mapped = all(col in saved_mapping for col in columns_list)
        if all_mapped:
            filtered_mapping = {col: saved_mapping[col] for col in columns_list}
            result = build_detection_result(columns_list, filtered_mapping)
            result["file_id"] = file_id
            result["source"] = "saved_mapping"
            result["message"] = "Mapping loaded from saved client profile — LLM skipped."
            return result
    # No saved mapping, find the uploaded file
    save_path = None
    file_ext = None
    for extension in ALLOWED_EXTENSIONS:
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{extension}")
        if os.path.exists(path):
            save_path = path
            file_ext = extension
            break
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")
    # Read file and extract first non-empty value per column as sample context
    try:
        df = read_file_to_df(save_path, file_ext)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
        sample_values = {}
        for col in columns_list:
            if col in df.columns:
                non_empty = df[col].dropna().replace("", float("nan")).dropna()
                sample_values[col] = str(non_empty.iloc[0]) if len(non_empty) > 0 else ""
            else:
                sample_values[col] = ""
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")

    # Run LLM detection with column names, sample values and fill rates
    try:
        mapping = detect_columns_with_llm(columns_list, sample_values, fill_rates_dict)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Detection failed: {str(e)}")
    result = build_detection_result(columns_list, mapping)
    result["file_id"] = file_id
    result["source"] = "llm_detection"
    return result

# Endpoint to save confirmed column mapping for a client
@app.post("/save-mapping")
async def save_mapping_endpoint(
    client_id: str = Form(...),
    file_type: str = Form(...),
    mapping: str = Form(...),
    confirmed_by: str = Form(None)
):
    try:
        mapping_dict = json.loads(mapping)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid mapping format.")
    if not mapping_dict:
        raise HTTPException(status_code=400, detail="Mapping cannot be empty. Please provide a valid mapping.")
    save_mapping(client_id, file_type, mapping_dict, confirmed_by)
    return {
        "client_id": client_id,
        "file_type": file_type,
        "columns_saved": len(mapping_dict),
        "message": f"Mapping saved successfully for client {client_id} and file type {file_type}."
    }

# Endpoint to retrieve saved mapping for a client
@app.get("/get-mapping/{client_id}")
async def get_mapping_endpoint(client_id: str, file_type: str = "general"):
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        return {
            "client_id": client_id,
            "file_type": file_type,
            "mapping": {},
            "message": "No saved mapping found for this client."
        }
    return {
        "client_id": client_id,
        "file_type": file_type,
        "mapping": mapping,
        "columns_mapped": len(mapping),
        "message": "Saved mapping retrieved successfully."
    }

# Endpoint to get upload history for a client
@app.get("/uploads/{client_id}")
async def get_uploads_endpoint(client_id: str):
    uploads = get_uploads(client_id)
    return {
        "client_id": client_id,
        "total_uploads": len(uploads),
        "uploads": uploads
    }

# Endpoint to clean the uploaded file using confirmed mapping
@app.post("/clean")
async def clean_file(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general")
) -> dict:
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        raise HTTPException(
            status_code=400,
            detail="No saved mapping found for this client. Please detect the columns and confirm the mapping first."
        )
    save_path = None
    file_ext = None
    for extension in ALLOWED_EXTENSIONS:
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{extension}")
        if os.path.exists(path):
            save_path = path
            file_ext = extension
            break
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")
    try:
        df = read_file_to_df(save_path, file_ext)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")
    try:
        cleaned_df, report = clean_dataframe(df, mapping)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cleaning failed: {str(e)}")
    return {
        "file_id": file_id,
        "client_id": client_id,
        "file_type": file_type,
        "cleaned_data": cleaned_df.fillna("").to_dict(orient="records"),
        "validation_report": report,
        "message": "File cleaned successfully."
    }

